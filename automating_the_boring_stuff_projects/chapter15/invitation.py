import docx
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

def preparing_invitation():
    doc = docx.Document()
    # get the guests' name from the file
    with open('guests.txt', 'r') as name_list:
        names = name_list.readlines()

        run = []
        for name in names:
            run.append(doc.add_paragraph().add_run('It would be a pleasure to have the company of'))
            run.append(doc.add_paragraph().add_run(name))
            run.append(doc.add_paragraph().add_run('at 11010 Memory Lane on the Evening of'))
            run.append(doc.add_paragraph().add_run('April 1st'))
            run.append(doc.add_paragraph().add_run('at 7 o\'clock'))

            # look for the last 5 runs
            for i in range(len(doc.paragraphs) - 5, len(doc.paragraphs)):
                # align center
                doc.paragraphs[i].paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER

                # do the styling
                font = run[i].font
                if i == len(doc.paragraphs) - 5 or i == len(doc.paragraphs) - 3 or i == len(doc.paragraphs) - 1:
                    font.name = 'Forte'
                else:
                    font.name = 'Arial'
                font.size = Pt(16)

            # make the guest name bold
            run[len(doc.paragraphs) - 4].bold = True

            # go to the next page for the next invitation
            run[len(doc.paragraphs) - 1].add_break(docx.enum.text.WD_BREAK.PAGE)

    doc.save('invitation.docx')

def main():
    preparing_invitation()

if __name__ == '__main__':
    main()
